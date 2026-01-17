#!/usr/bin/env python3
"""
영한 특허번역 마크다운 → 워드 변환 스크립트

사용법:
    python convert-to-docx.py <input_markdown> [output_docx]

예시:
    python convert-to-docx.py translation-final.md
    python convert-to-docx.py translation-final.md output.docx

포맷 설정:
    - 폰트: 바탕체 (Batang)
    - 폰트 크기: 12pt
    - 문단 정렬: 양쪽 정렬 (Justify)
    - 줄 간격: 1.5줄
"""

import sys
import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx 패키지가 필요합니다.")
    print("설치: pip install python-docx")
    sys.exit(1)


def set_korean_font(run, font_name='Batang', font_size=12):
    """한글 폰트 설정"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    # 한글 폰트를 위한 추가 설정
    r = run._element
    rFonts = r.find(qn('w:rPr'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rFonts)
    fonts = rFonts.find(qn('w:rFonts'))
    if fonts is None:
        fonts = rFonts.makeelement(qn('w:rFonts'), {})
        rFonts.append(fonts)
    fonts.set(qn('w:eastAsia'), font_name)


def parse_markdown(md_content):
    """마크다운 파싱"""
    lines = md_content.split('\n')
    elements = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # 빈 줄
        if not line.strip():
            elements.append({'type': 'blank'})
            i += 1
            continue

        # 헤딩 (# ~ ######)
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            elements.append({'type': 'heading', 'level': level, 'text': text})
            i += 1
            continue

        # 수평선 (---, ***, ___)
        if re.match(r'^[-*_]{3,}\s*$', line):
            elements.append({'type': 'hr'})
            i += 1
            continue

        # 테이블
        if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            elements.append({'type': 'table', 'lines': table_lines})
            continue

        # 코드 블록
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append({'type': 'code', 'lines': code_lines})
            i += 1
            continue

        # 리스트 아이템
        list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', line)
        if list_match:
            indent = len(list_match.group(1))
            marker = list_match.group(2)
            text = list_match.group(3)
            is_ordered = marker[0].isdigit()
            elements.append({'type': 'list', 'indent': indent, 'ordered': is_ordered, 'text': text})
            i += 1
            continue

        # 일반 텍스트
        elements.append({'type': 'paragraph', 'text': line})
        i += 1

    return elements


def clean_markdown_formatting(text):
    """마크다운 인라인 서식 제거 (굵게, 기울임 등)"""
    # **bold** 또는 __bold__
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # *italic* 또는 _italic_
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # `code`
    text = re.sub(r'`(.+?)`', r'\1', text)
    # [link](url)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def parse_table(table_lines):
    """테이블 파싱"""
    rows = []
    for line in table_lines:
        # 구분선 (|---|---|) 건너뛰기
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    return rows


def convert_to_docx(md_path, docx_path=None, template_path=None):
    """마크다운을 워드로 변환"""

    # 입력 파일 확인
    md_path = Path(md_path)
    if not md_path.exists():
        print(f"Error: 파일을 찾을 수 없습니다: {md_path}")
        return False

    # 출력 경로 설정
    if docx_path is None:
        docx_path = md_path.with_suffix('.docx')
    else:
        docx_path = Path(docx_path)

    # 마크다운 읽기
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
    except FileNotFoundError:
        print(f"Error: 마크다운 파일을 찾을 수 없습니다: {md_path}")
        return False
    except PermissionError:
        print(f"Error: 마크다운 파일 읽기 권한이 없습니다: {md_path}")
        return False
    except UnicodeDecodeError:
        print(f"Error: 마크다운 파일 인코딩 오류 (UTF-8 필요): {md_path}")
        return False
    except Exception as e:
        print(f"Error: 마크다운 파일 읽기 실패: {md_path} - {str(e)}")
        return False

    # 템플릿 사용 또는 새 문서 생성
    if template_path and Path(template_path).exists():
        doc = Document(template_path)
    else:
        doc = Document()

        # 기본 스타일 설정
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Batang'
        font.size = Pt(12)

        paragraph_format = style.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.line_spacing = 1.5

        # 한글 폰트 설정
        style_element = style._element
        rPr = style_element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), 'Batang')

    # 마크다운 파싱
    elements = parse_markdown(md_content)

    # 워드 문서 생성
    for elem in elements:
        if elem['type'] == 'blank':
            # 빈 줄은 건너뛰기 (연속 빈 줄 방지)
            continue

        elif elem['type'] == 'heading':
            level = elem['level']
            text = clean_markdown_formatting(elem['text'])

            # 헤딩 스타일 적용
            if level == 1:
                p = doc.add_heading(text, level=1)
            elif level == 2:
                p = doc.add_heading(text, level=2)
            else:
                p = doc.add_heading(text, level=min(level, 4))

            # 헤딩에도 바탕체 적용
            for run in p.runs:
                set_korean_font(run, 'Batang', 14 if level == 1 else 12)

        elif elem['type'] == 'paragraph':
            text = clean_markdown_formatting(elem['text'])
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                set_korean_font(run)

        elif elem['type'] == 'list':
            text = clean_markdown_formatting(elem['text'])
            indent_level = elem['indent'] // 2

            if elem['ordered']:
                p = doc.add_paragraph(text, style='List Number')
            else:
                p = doc.add_paragraph(text, style='List Bullet')

            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            for run in p.runs:
                set_korean_font(run)

        elif elem['type'] == 'table':
            rows = parse_table(elem['lines'])
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'

                for i, row in enumerate(rows):
                    for j, cell_text in enumerate(row):
                        cell = table.cell(i, j)
                        cell.text = clean_markdown_formatting(cell_text)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                set_korean_font(run, font_size=10)

                # 테이블 후 빈 줄 추가
                doc.add_paragraph()

        elif elem['type'] == 'code':
            # 코드 블록은 들여쓰기된 단락으로
            for line in elem['lines']:
                p = doc.add_paragraph(line)
                p.paragraph_format.left_indent = Inches(0.5)
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)

        elif elem['type'] == 'hr':
            # 수평선은 빈 줄로 대체
            doc.add_paragraph()

    # 저장
    try:
        doc.save(str(docx_path))
    except PermissionError:
        print(f"Error: 워드 파일 쓰기 권한이 없습니다: {docx_path}")
        return False
    except OSError as e:
        print(f"Error: 워드 파일 쓰기 실패: {docx_path} - {str(e)}")
        return False
    except Exception as e:
        print(f"Error: 예상치 못한 쓰기 오류: {docx_path} - {str(e)}")
        return False

    print(f"변환 완료: {docx_path}")
    return True


def merge_sections(project_dir, output_file='translation-final.md'):
    """섹션 파일들을 하나로 병합"""

    project_dir = Path(project_dir)
    sections_dir = project_dir / 'sections'

    if not sections_dir.exists():
        print(f"Error: sections 폴더를 찾을 수 없습니다: {sections_dir}")
        return None

    # 섹션 파일 정렬 (section-01, section-02, ...)
    section_files = sorted(sections_dir.glob('section-*.md'))

    if not section_files:
        print(f"Error: 섹션 파일을 찾을 수 없습니다: {sections_dir}")
        return None

    # 병합
    merged_content = []
    for section_file in section_files:
        with open(section_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # 메타데이터 제거 (# 섹션 N: 이후 첫 --- 까지)
        lines = content.split('\n')
        start_idx = 0
        hr_count = 0

        for i, line in enumerate(lines):
            if line.strip() == '---':
                hr_count += 1
                if hr_count == 1:
                    start_idx = i + 1
                    break

        # 번역문 섹션만 추출
        in_translation = False
        translation_lines = []

        for line in lines[start_idx:]:
            if '## 번역문' in line:
                in_translation = True
                continue
            if in_translation:
                if line.startswith('## ') or line.strip() == '---':
                    break
                translation_lines.append(line)

        if translation_lines:
            merged_content.append('\n'.join(translation_lines))
        else:
            # 번역문 섹션이 없으면 전체 내용 사용
            merged_content.append('\n'.join(lines[start_idx:]))

    # 저장
    output_path = project_dir / output_file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n\n---\n\n'.join(merged_content))

    print(f"병합 완료: {output_path}")
    return output_path


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    # 프로젝트 디렉토리인 경우 병합 후 변환
    if os.path.isdir(input_path):
        merged_file = merge_sections(input_path)
        if merged_file:
            convert_to_docx(merged_file)
    else:
        # 마크다운 파일 직접 변환
        convert_to_docx(input_path, output_path)


if __name__ == '__main__':
    main()
